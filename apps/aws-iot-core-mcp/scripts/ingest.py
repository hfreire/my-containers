#!/usr/bin/env python3
"""
Ingest Espen knowledge from a private GitHub repo into Qdrant.

Sources:
  - docs/configs/<manufacturer>/<model>.json   → register mappings
  - docs/configs/<manufacturer>/spec/*.pdf     → manufacturer Modbus specs (PDF)
  - docs/configs/<manufacturer>/spec/*.docx    → manufacturer Modbus specs (DOCX)
  - docs/configs/<manufacturer>/spec/*.xlsx    → manufacturer Modbus specs (XLSX)
  - docs/schemas/<shadow-state>.json5          → shadow state JSON schemas

Usage:
  GITHUB_TOKEN=ghp_xxx \
  QDRANT_URL=http://qdrant.databases:6333 \
  QDRANT_API_KEY=xxx \
  EMBEDDING_URL=http://litellm.default:4000/v1/embeddings \
  EMBEDDING_API_KEY=xxx \
  python3 ingest.py
"""

import hashlib
import json
import os
import sys
import time
from typing import Any

import requests

# Config
GITHUB_TOKEN = os.environ["GITHUB_TOKEN"]
GITHUB_REPO = os.environ.get("GITHUB_REPO", "feverenergy/espen")
GITHUB_BRANCH = os.environ.get("GITHUB_BRANCH", "main")

QDRANT_URL = os.environ.get("QDRANT_URL", "http://qdrant.databases:6333")
QDRANT_API_KEY = os.environ.get("QDRANT_API_KEY", "")
QDRANT_COLLECTION = os.environ.get("QDRANT_COLLECTION", "espen-knowledge")

EMBEDDING_URL = os.environ.get(
    "EMBEDDING_URL", "http://litellm.default:4000/v1/embeddings"
)
EMBEDDING_API_KEY = os.environ.get("EMBEDDING_API_KEY", "")
EMBEDDING_MODEL = os.environ.get("EMBEDDING_MODEL", "text-embedding-3-small")

CHUNK_MAX_TOKENS = 600
VERIFY_SSL = os.environ.get("VERIFY_SSL", "true").lower() != "false"

github_headers = {"Authorization": f"token {GITHUB_TOKEN}"}
qdrant_headers = {"Content-Type": "application/json"}
if QDRANT_API_KEY:
    qdrant_headers["api-key"] = QDRANT_API_KEY


# ── Helpers ─────────────────────────────────────────────────────────


def github_get(path: str) -> Any:
    url = f"https://api.github.com/repos/{GITHUB_REPO}/{path}"
    r = requests.get(url, headers=github_headers, verify=VERIFY_SSL)
    r.raise_for_status()
    return r.json()


def github_raw(file_path: str) -> bytes:
    url = f"https://raw.githubusercontent.com/{GITHUB_REPO}/{GITHUB_BRANCH}/{file_path}"
    r = requests.get(url, headers=github_headers, verify=VERIFY_SSL)
    r.raise_for_status()
    return r.content


def get_tree() -> list[dict]:
    return github_get(f"git/trees/{GITHUB_BRANCH}?recursive=1")["tree"]


def get_embedding(text: str) -> list[float]:
    headers = {"Content-Type": "application/json"}
    if EMBEDDING_API_KEY:
        headers["Authorization"] = f"Bearer {EMBEDDING_API_KEY}"
    r = requests.post(
        EMBEDDING_URL,
        headers=headers,
        json={"model": EMBEDDING_MODEL, "input": [text]},
        verify=VERIFY_SSL,
    )
    r.raise_for_status()
    return r.json()["data"][0]["embedding"]


def upsert_point(point_id: str, vector: list[float], payload: dict) -> None:
    r = requests.put(
        f"{QDRANT_URL}/collections/{QDRANT_COLLECTION}/points",
        headers=qdrant_headers,
        json={"points": [{"id": point_id, "vector": vector, "payload": payload}]},
        verify=VERIFY_SSL,
    )
    r.raise_for_status()


def chunk_id(text: str) -> str:
    return hashlib.md5(text.encode()).hexdigest()


def rough_token_count(text: str) -> int:
    return len(text.split())


def chunk_text(text: str, max_tokens: int = CHUNK_MAX_TOKENS) -> list[str]:
    """Split text into chunks at paragraph boundaries."""
    paragraphs = text.split("\n\n")
    chunks = []
    current = ""

    for para in paragraphs:
        if rough_token_count(current + "\n\n" + para) > max_tokens and current:
            chunks.append(current.strip())
            sentences = current.strip().split(". ")
            overlap = sentences[-1] if len(sentences) > 1 else ""
            current = overlap + "\n\n" + para if overlap else para
        else:
            current = current + "\n\n" + para if current else para

    if current.strip():
        chunks.append(current.strip())

    return chunks


def ingest_chunks(chunks: list[str], file_path: str, doc_type: str, metadata: dict) -> int:
    """Embed and upsert a list of text chunks. Returns count ingested."""
    count = 0
    for i, chunk in enumerate(chunks):
        if rough_token_count(chunk) < 20:
            continue

        cid = chunk_id(f"{file_path}_{i}")
        payload = {
            "text": chunk,
            "source": file_path,
            "type": doc_type,
            "chunk_index": i,
            **metadata,
        }

        try:
            vector = get_embedding(chunk)
            upsert_point(cid, vector, payload)
            count += 1
        except Exception as e:
            print(f"    Chunk {i} error: {e}")

        if i % 10 == 9:
            time.sleep(0.5)

    return count


# ── Text extractors ─────────────────────────────────────────────────


def extract_pdf_text(data: bytes) -> str:
    try:
        import fitz
    except ImportError:
        print("    pymupdf not installed — pip install pymupdf")
        return ""
    doc = fitz.open(stream=data, filetype="pdf")
    text = "\n\n".join(page.get_text() for page in doc)
    doc.close()
    return text.strip()


def extract_docx_text(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError:
        print("    python-docx not installed — pip install python-docx")
        return ""
    from io import BytesIO

    doc = Document(BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    return "\n\n".join(paragraphs)


def extract_xlsx_text(data: bytes) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError:
        print("    openpyxl not installed — pip install openpyxl")
        return ""
    from io import BytesIO

    wb = load_workbook(BytesIO(data), read_only=True, data_only=True)
    sections = []

    for sheet in wb.sheetnames:
        ws = wb[sheet]
        rows = []
        header = None
        for row in ws.iter_rows(values_only=True):
            cells = [str(c).strip() if c is not None else "" for c in row]
            if not any(cells):
                continue
            if header is None:
                header = cells
                rows.append(" | ".join(cells))
            else:
                # Format as "header: value" pairs for better semantic search
                pairs = []
                for h, v in zip(header, cells):
                    if v and h:
                        pairs.append(f"{h}: {v}")
                if pairs:
                    rows.append(", ".join(pairs))

        if rows:
            sections.append(f"Sheet: {sheet}\n" + "\n".join(rows))

    wb.close()
    return "\n\n".join(sections)


# ── Ingestors ───────────────────────────────────────────────────────


def ingest_register_configs(tree: list[dict]) -> int:
    """Ingest JSON register configs from docs/configs/<manufacturer>/<model>.json"""
    json_files = [
        item["path"]
        for item in tree
        if item["path"].startswith("docs/configs/")
        and item["path"].endswith(".json")
        and "/spec/" not in item["path"]
    ]

    count = 0
    for file_path in json_files:
        print(f"  Processing: {file_path}")
        try:
            content = json.loads(github_raw(file_path))
        except Exception as e:
            print(f"    Skip (parse error): {e}")
            continue

        manufacturer = content.get("manufacturer", "unknown")
        model = content.get("model", "unknown")
        registers = content.get("registers", [])

        if not registers:
            print(f"    Skip (no registers)")
            continue

        for i in range(0, len(registers), 10):
            batch = registers[i : i + 10]
            lines = [
                f"Manufacturer: {manufacturer} | Model: {model}",
                f"Register mapping (registers {i+1}-{i+len(batch)}):",
                "",
            ]
            for reg in batch:
                name = reg.get("display_name", "unknown")
                addr = reg.get("address", "?")
                reg_type = reg.get("register_type", "?")
                val_type = reg.get("value_type", "?")
                size = reg.get("address_size", 1)
                line = f"- {name}: register {addr} ({reg_type}), type={val_type}, size={size}"
                if "scale" in reg:
                    line += f", scale={reg['scale']}"
                if "unit" in reg:
                    line += f", unit={reg['unit']}"
                lines.append(line)

            chunk_text_str = "\n".join(lines)
            cid = chunk_id(f"{file_path}_{i}")

            try:
                vector = get_embedding(chunk_text_str)
                upsert_point(
                    cid,
                    vector,
                    {
                        "text": chunk_text_str,
                        "source": file_path,
                        "type": "register_mapping",
                        "manufacturer": manufacturer,
                        "model": model,
                        "chunk_index": i // 10,
                    },
                )
                count += 1
                print(f"    Chunk {i//10}: {len(batch)} registers")
            except Exception as e:
                print(f"    Error: {e}")

        time.sleep(0.2)

    return count


def ingest_manufacturer_specs(tree: list[dict]) -> int:
    """Ingest specs from docs/configs/<manufacturer>/spec/ (PDF, DOCX, XLSX)"""
    spec_files = [
        item["path"]
        for item in tree
        if item["path"].startswith("docs/configs/")
        and "/spec/" in item["path"]
        and any(item["path"].lower().endswith(ext) for ext in [".pdf", ".docx", ".xlsx"])
    ]

    count = 0
    for file_path in spec_files:
        parts = file_path.split("/")
        manufacturer = parts[2] if len(parts) > 2 else "unknown"
        filename = parts[-1]
        ext = filename.rsplit(".", 1)[-1].lower()

        print(f"  Processing: {file_path}")
        try:
            data = github_raw(file_path)
        except Exception as e:
            print(f"    Skip (download error): {e}")
            continue

        if ext == "pdf":
            text = extract_pdf_text(data)
        elif ext == "docx":
            text = extract_docx_text(data)
        elif ext == "xlsx":
            text = extract_xlsx_text(data)
        else:
            continue

        if not text or len(text) < 100:
            print(f"    Skip (no text extracted)")
            continue

        chunks = chunk_text(text)
        print(f"    Extracted {len(chunks)} chunks from {rough_token_count(text)} tokens")

        prefixed_chunks = [
            f"Manufacturer: {manufacturer}\nSource: {filename}\n\n{chunk}"
            for chunk in chunks
        ]

        ingested = ingest_chunks(
            prefixed_chunks,
            file_path,
            "inverter_spec",
            {"manufacturer": manufacturer, "filename": filename},
        )
        count += ingested
        time.sleep(0.2)

    return count


def ingest_shadow_schemas(tree: list[dict]) -> int:
    """Ingest shadow state schemas from docs/schemas/<shadow-state>.json5"""
    schema_files = [
        item["path"]
        for item in tree
        if item["path"].startswith("docs/schemas/")
        and item["path"].endswith(".json5")
    ]

    count = 0
    for file_path in schema_files:
        print(f"  Processing: {file_path}")
        filename = file_path.split("/")[-1]
        shadow_name = filename.rsplit(".", 1)[0]

        try:
            raw = github_raw(file_path)
        except Exception as e:
            print(f"    Skip (download error): {e}")
            continue

        # Parse JSON5
        try:
            import json5

            schema = json5.loads(raw.decode("utf-8"))
        except ImportError:
            print("    json5 not installed — pip install json5")
            continue
        except Exception as e:
            print(f"    Skip (parse error): {e}")
            continue

        # Extract field documentation from schema properties
        fields = extract_schema_fields(schema, shadow_name)
        print(f"    Extracted {len(fields)} field descriptions")

        # Group fields into batches of ~10 per chunk for sufficient context
        chunks = []
        for i in range(0, len(fields), 10):
            batch = fields[i : i + 10]
            header = f"Shadow state: {shadow_name}\nField documentation:\n\n"
            chunks.append(header + "\n\n".join(batch))

        ingested = ingest_chunks(
            chunks,
            file_path,
            "shadow_schema",
            {"shadow_name": shadow_name},
        )
        count += ingested
        time.sleep(0.2)

    return count


def extract_schema_fields(schema: Any, shadow_name: str, prefix: str = "") -> list[str]:
    """Recursively extract field descriptions from a JSON Schema."""
    chunks = []

    if not isinstance(schema, dict):
        return chunks

    properties = schema.get("properties", {})
    for field_name, field_def in properties.items():
        if not isinstance(field_def, dict):
            continue

        full_path = f"{prefix}.{field_name}" if prefix else field_name
        lines = [f"Shadow: {shadow_name} | Field: {full_path}"]

        if "type" in field_def:
            lines.append(f"Type: {field_def['type']}")
        if "description" in field_def:
            lines.append(f"Description: {field_def['description']}")
        if "enum" in field_def:
            lines.append(f"Values: {', '.join(str(v) for v in field_def['enum'])}")
        if "default" in field_def:
            lines.append(f"Default: {field_def['default']}")
        if "minimum" in field_def or "maximum" in field_def:
            range_parts = []
            if "minimum" in field_def:
                range_parts.append(f"min={field_def['minimum']}")
            if "maximum" in field_def:
                range_parts.append(f"max={field_def['maximum']}")
            lines.append(f"Range: {', '.join(range_parts)}")

        # Only create a chunk if there's meaningful info beyond just the name
        if len(lines) > 1:
            chunks.append("\n".join(lines))

        # Recurse into nested objects
        if field_def.get("type") == "object":
            chunks.extend(extract_schema_fields(field_def, shadow_name, full_path))

        # Recurse into array items
        if field_def.get("type") == "array" and "items" in field_def:
            chunks.extend(
                extract_schema_fields(field_def["items"], shadow_name, f"{full_path}[]")
            )

    return chunks


# ── Main ────────────────────────────────────────────────────────────


def main():
    print("=" * 60)
    print("Espen Knowledge Ingestion")
    print("=" * 60)
    print(f"Repo: {GITHUB_REPO} ({GITHUB_BRANCH})")
    print(f"Qdrant: {QDRANT_URL} / {QDRANT_COLLECTION}")
    print(f"Embedding: {EMBEDDING_URL} ({EMBEDDING_MODEL})")
    print()

    tree = get_tree()

    print("── Register Mappings (JSON) ──")
    reg_count = ingest_register_configs(tree)
    print(f"  → {reg_count} chunks ingested")
    print()

    print("── Manufacturer Specs (PDF/DOCX/XLSX) ──")
    spec_count = ingest_manufacturer_specs(tree)
    print(f"  → {spec_count} chunks ingested")
    print()

    print("── Shadow State Schemas (JSON5) ──")
    schema_count = ingest_shadow_schemas(tree)
    print(f"  → {schema_count} chunks ingested")
    print()

    total = reg_count + spec_count + schema_count
    print(f"Done. Total: {total} chunks ingested into {QDRANT_COLLECTION}.")

    if total == 0:
        sys.exit(1)


if __name__ == "__main__":
    main()
